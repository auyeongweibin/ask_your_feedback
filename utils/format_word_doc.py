from docx import Document, enum
from docx.shared import Pt

def format_word_doc(data, school:str, name:str, type:str):
    document = Document()

    run = document.add_paragraph().add_run(type)
    run.font.name = 'Open Sans'
    run.font.size = Pt(12)

    if type == 'Original w/o Negative':
        for line in data:
            if line[0].isdigit():
                if school == 'SMU':
                    run = document.add_paragraph().add_run(line.split('\n')[1])
                    run.font.name = 'Open Sans'
                    run.font.size = Pt(7)
                    
                elif school == 'UW':
                    run = document.add_paragraph().add_run(''.join(line.split('. ')[1:]))
                    run.font.name = 'Open Sans'
                    run.font.size = Pt(7)
                    
            else:
                run = document.add_paragraph().add_run('\n' + line)
                run.font.name = 'Open Sans'
                run.font.size = Pt(7)
                run.bold = True
                run.font.highlight_color = enum.text.WD_COLOR.GRAY_50
    elif type == 'Original with Actionable Feedback':
        pass
    elif type == 'ET&L':
        pass
    elif type == 'Letter Format Sandwich':
        pass
    
    document.save(name)
