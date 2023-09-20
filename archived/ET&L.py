import streamlit as st
import os
from utils.classify import classify
from utils.extract import extract
from classes import categories
from docx import Document
from docx.shared import Pt

st.title('ET&L Codes')
st.markdown("## Hi Instructor, let me summarise your feedback for you 🎈")

file = st.file_uploader('Upload your feedback!')
if file:
    with open(os.path.join("./uploads/", file.name),"wb") as f:
        f.write(file.getbuffer())
        st.success("Saved File")
    
    qualitative, school = extract(file.name)

    qualitative = list(filter(lambda x: x.lower() not in ['nil', 'none', 'na', ''], qualitative))

    st.success("Extracted Qualitative Feedback")

    if len(qualitative) == 0:
        st.write('No Qualitative Feedback Found')
    else:
        coded = classify(qualitative, categories)

        st.success("Coded Qualitative Feedback")

        document = Document()

        run = document.add_paragraph().add_run('ET&L')
        run.font.name = 'Open Sans'
        run.font.size = Pt(12)

        for code in coded:
            run = document.add_paragraph().add_run(code)
            run.font.name = 'Open Sans'
            run.font.size = Pt(7)
            run.font.bold = True

            for line in coded[code]:
                run = document.add_paragraph(style='List Bullet').add_run(line)
                run.font.name = 'Open Sans'
                run.font.size = Pt(7)

        document.save(os.path.join("./processed_pdfs/", file.name.replace('.pdf', '_ET&L.docx')))

        with open(os.path.join("./processed_pdfs/", file.name.replace('.pdf', '_ET&L.docx')),"rb") as f:
            st.download_button(
                label='Download ET&L',
                data=f,
                file_name=file.name.replace('.pdf', '_ET&L.docx')
            )

        

    