import streamlit as st
import os
from utils.extract import extract_with_questions
from utils.generate import generate
from docx import Document, enum
from docx.shared import Pt

st.title('Actionable Feedback')
st.markdown("## Hi Instructor, let me summarise your feedback for you 🎈")

file = st.file_uploader('Upload your feedback!')
if file:
    with open(os.path.join("./uploads/", file.name),"wb") as f:
        f.write(file.getbuffer())
        st.success("Saved File")
    
    qualitative, school = extract_with_questions(file.name)

    qualitative = list(filter(lambda x: len(x.split('\n'))==1 or x.split('\n')[1].lower() not in ['nil', 'nil.', 'none', 'na', ''], qualitative))

    st.success("Extracted Qualitative Feedback")

    if len(qualitative) == 0:
        st.write('No Qualitative Feedback Found')
    else:
        questions = list(filter(lambda x: len(x) and not x[0].isdigit(), qualitative))

        format = '\n'.join([f"{i+1}. 'actionable feedback for question {i+1}'" for i in range(len(questions))])

        prompt = f'''
            For each question and set of associated questions in parenthesis, suggest actionable feedback for each question, and output it in the format in triple backticks:
            ({qualitative})

            ```
            {format}
            ```
        '''

        # response = generate(prompt, 'bard')
        response = generate(prompt, 'gpt-3.5-turbo')

        result = list(filter(lambda x: x != '', response.split('\n')))
        # result = list(filter(lambda x: not x[0].isdigit(), result))
        result = list(filter(lambda x: x[0].isdigit(), result))

        document = Document()

        run = document.add_paragraph().add_run('Original with Actionable Feedback')
        run.font.name = 'Open Sans'
        run.font.size = Pt(12)

        question_number = 0

        for line in qualitative:
            if line[0].isdigit():
                if school == 'SMU':
                    run = document.add_paragraph().add_run(line.split('\n')[1])
                    run.font.name = 'Open Sans'
                    run.font.size = Pt(7)
                    
                elif school == 'UW':
                    run = document.add_paragraph().add_run(''.join(line.split('. ')[1:]))
                    run.font.name = 'Open Sans'
                    run.font.size = Pt(7)
                    
            else:
                run = document.add_paragraph().add_run('\n' + line)
                run.font.name = 'Open Sans'
                run.font.size = Pt(7)
                run.bold = True

                # run = document.add_paragraph().add_run(result[question_number])
                run = document.add_paragraph().add_run(result[question_number].split('. ')[1])
                run.font.name = 'Open Sans'
                run.font.size = Pt(7)
                run.font.highlight_color = enum.text.WD_COLOR.YELLOW
                question_number += 1
        
        document.save(os.path.join("./processed_pdfs/", file.name.replace('.pdf', '_Actionable_Feedback.docx')))

        with open(os.path.join("./processed_pdfs/", file.name.replace('.pdf', '_Actionable_Feedback.docx')),"rb") as f:
            st.download_button(
                label='Download Actionable Feedback',
                data=f,
                file_name=file.name.replace('.pdf', '_Actionable_Feedback.docx')
            )
