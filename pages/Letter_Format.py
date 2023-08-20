import streamlit as st
import os
from utils.generate import generate
from utils.extract import extract
from docx import Document
from docx.shared import Pt

st.title('Letter Format')
st.markdown("## Hi Instructor, let me summarise your feedback for you 🎈")

file = st.file_uploader('Upload your feedback!')
if file:
    with open(os.path.join("./uploads/", file.name),"wb") as f:
        f.write(file.getbuffer())
        st.success("Saved File")
    
    qualitative, school = extract(file.name)

    st.success("Extracted Qualitative Feedback")

    if len(qualitative) == 0:
        st.write('No Qualitative Feedback Found')
    else:
        prompt = f'''
            Using the data in parenthesis, write a letter to summarise the feedback in a constructive manner, a format is given in triple backticks.
            If there is any negative feedback, use the sandwich method, i.e. position negative feedback between 2 other pieces of positive feedback, e.g. 'positive feedback', 'negative feedback'. However, 'positive feedback'

            ({qualitative})

            ```
            Subject: Positive Feedback and Constructive Suggestions for Improvement - XXX Class

            Dear Instructor,

            'insert feedback here'

            Best regards,

            Ken-bot
            ```
        '''

        response = generate(prompt, 'gpt-3.5-turbo' if school == 'SMU' else 'gpt-3.5-turbo-16k')
        
        document = Document()

        run = document.add_paragraph().add_run('Letter Format Sandwich')
        run.font.name = 'Open Sans'
        run.font.size = Pt(12)

        run = document.add_paragraph().add_run(response)
        run.font.name = 'Open Sans'
        run.font.size = Pt(7)

        document.save(os.path.join("./processed_pdfs/", file.name.replace('.pdf', '_Letter_Format_Sandwich.docx')))

        with open(os.path.join("./processed_pdfs/", file.name.replace('.pdf', '_Letter_Format_Sandwich.docx')),"rb") as f:
            st.download_button(
                label='Download Letter Format Sandwich',
                data=f,
                file_name=file.name.replace('.pdf', '_Letter_Format_Sandwich.docx')
            )