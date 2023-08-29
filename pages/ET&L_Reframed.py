import streamlit as st
import os
from utils.classify import classify
from classes import categories
from utils.generate import generate
from utils.extract import extract
from docx import Document
from docx.shared import Pt


st.title('ET&L Reframed')
st.markdown("## Hi Instructor, let me summarise your feedback for you 🎈")

file = st.file_uploader('Upload your feedback!')
if file:
    with open(os.path.join("./uploads/", file.name),"wb") as f:
        f.write(file.getbuffer())
        st.success("Saved File")
    
    qualitative, school = extract(file.name)

    qualitative = list(filter(lambda x: x.lower() not in ['nil', 'none', 'na', ''], qualitative))

    st.success("Extracted Qualitative Feedback")

    if len(qualitative) == 0:
        st.write('No Qualitative Feedback Found')
    else:
        coded = classify(qualitative, categories)
        st.success("Coded Qualitative Feedback")

        prompt = f'''
            Using the data in parenthesis, generate a summary of the feedback in each class and output it in the format in triple backticks
            Rephrase any negative feedback or comments in a positive and constructive manner

            ({coded})

            ```
            for each code/category:
            Code/Category: 'summary of feedback in the code/category'
            ```
        '''

        # response = generate(prompt, 'gpt-3.5-turbo' if school == 'SMU' else 'bard')
        response = generate(prompt, 'gpt-3.5-turbo')

        st.write(response)

        result = list(filter(lambda x: x not in ['', '```'], response.split('\n')))

        document = Document()

        run = document.add_paragraph().add_run('ET&L Reframed')
        run.font.name = 'Open Sans'
        run.font.size = Pt(12)

        feedback = ''

        for line in result:
            # if line.startswith('Code/Category'):
            if not line.startswith('- '):
                if feedback != '':
                    run = document.add_paragraph(style='List Bullet').add_run(feedback)
                    run.font.name = 'Open Sans'
                    run.font.size = Pt(7)

                    feedback = ''

                run = document.add_paragraph().add_run(line[:-1])
                run.font.name = 'Open Sans'
                run.font.size = Pt(7)
                run.font.bold = True
            else:
                feedback += line[2:]
        
        if feedback != '':
            run = document.add_paragraph(style='List Bullet').add_run(feedback[1:])
            run.font.name = 'Open Sans'
            run.font.size = Pt(7)
            
        document.save(os.path.join("./processed_pdfs/", file.name.replace('.pdf', '_ET&L_Reframed.docx')))

        with open(os.path.join("./processed_pdfs/", file.name.replace('.pdf', '_ET&L_Reframed.docx')),"rb") as f:
            st.download_button(
                label='Download ET&L Reframed',
                data=f,
                file_name=file.name.replace('.pdf', '_ET&L_Reframed.docx')
            )

        

    