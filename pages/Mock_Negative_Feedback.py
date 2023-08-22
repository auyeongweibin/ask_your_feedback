import streamlit as st
import os
from utils.extract import extract_with_questions
from utils.generate import generate
from docx import Document, enum
from docx.shared import Pt, RGBColor


st.title('Mock Negative Feedback')
st.markdown("## Hi Instructor, let me add some negative feedback for you 🎈")

file = st.file_uploader('Upload your feedback!')
if file:
    with open(os.path.join("./uploads/", file.name),"wb") as f:
        f.write(file.getbuffer())
        st.success("Saved File")

    qualitative, school = extract_with_questions(file.name)

    qualitative = list(filter(lambda x: len(x.split('\n'))==1 or x.split('\n')[1].lower() not in ['nil', 'none', 'na', ''], qualitative))

    st.success("Extracted Qualitative Feedback")

    questions = list(filter(lambda x: len(x) and not x[0].isdigit(), qualitative))

    format = '\n'.join([f"{(i*2)+1}. 'first negative response to the question {i+1}'\n{(i*2)+2}. 'second negative response to the question {i+1}'" for i in range(len(questions))])

    if len(questions) == 0:
        st.write('No Qualitative Feedback Found')
    else:
        # TODO: Add negative feedback by question
        prompt = f'''
            generate and add 2 negative responses for each questions in parenthesis and output them in the format below
            ({questions})

            ```
            {format}
            ```
        '''

        result = generate(prompt, 'gpt-3.5-turbo')
        
        st.success('Generated Negative Feedback')

        st.markdown('#### Negative Feedback: ')

        result = list(filter(lambda x: x != '', result.split('\n')))

        question_number = 0

        document = Document()

        run = document.add_paragraph().add_run('Original with Mocked Negative Feedback')
        run.font.name = 'Open Sans'
        run.font.size = Pt(12)

        for line in qualitative:
            if line[0].isdigit():
                if school == 'SMU':
                    run = document.add_paragraph().add_run(line.split('\n')[1])
                    run.font.name = 'Open Sans'
                    run.font.size = Pt(7)
                    
                elif school == 'UW':
                    run = document.add_paragraph().add_run(''.join(line.split('. ')[1:]))
                    run.font.name = 'Open Sans'
                    run.font.size = Pt(7)
                    
            else:
                run = document.add_paragraph().add_run('\n' + line)
                run.font.name = 'Open Sans'
                run.font.size = Pt(7)
                run.bold = True

                run = document.add_paragraph().add_run(result[question_number][4:-1]+'\n\n'+result[question_number+1][4:-1])
                run.font.name = 'Open Sans'
                run.font.size = Pt(7)
                run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)
                question_number += 2
        
        document.save(os.path.join("./processed_pdfs/", file.name.replace('.pdf', '_Mock_Negative.docx')))

        with open(os.path.join("./processed_pdfs/", file.name.replace('.pdf', '_Mock_Negative.docx')),"rb") as f:
            st.download_button(
                label='Download Mock',
                data=f,
                file_name=file.name.replace('.pdf', '_Mock_Negative.docx')
            )

        